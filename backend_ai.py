import os
import json
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pydantic import BaseModel, Field
from typing import List, Set
from pypdf import PdfReader
from dotenv import load_dotenv

load_dotenv()

# --- DATA STRUCTURES ---
class ATSEvaluation(BaseModel):
    score: int
    missing_keywords: List[str]
    suggestions: List[str]

class ExperienceItem(BaseModel):
    role: str = ""
    company: str = ""
    duration: str = ""
    location: str = ""
    summary_input: str = ""
    tech_stack: str = ""
    enhanced_bullets: List[str] = Field(default_factory=list)

class ResumeData(BaseModel):
    full_name: str = ""
    contact_info: str = ""
    education: List[dict] = Field(default_factory=list)
    skills: dict = Field(default_factory=dict)
    experience: List[ExperienceItem] = Field(default_factory=list)

# --- UTILITIES ---
def extract_text_from_pdf(file_stream):
    reader = PdfReader(file_stream)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

# --- AI AGENT ---
class ResumeAgent:
    def __init__(self):
        api_key = os.getenv("OPENROUTER_API_KEY")
        if not api_key:
            raise ValueError("OPENROUTER_API_KEY missing")
            
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model_name = "google/gemini-1.5-flash" 

    def _call_llm(self, system_prompt: str, user_prompt: str):
        try:
            completion = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                response_format={"type": "json_object"} 
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"API Error: {e}")
            return None

    def parse_resume_text(self, raw_text: str) -> ResumeData:
        system_prompt = """
        You are a Resume Parser. Extract data into this exact JSON structure:
        {
            "full_name": "Name",
            "contact_info": "Phone | Email | LinkedIn",
            "education": [{"school": "", "degree": "", "year": ""}],
            "skills": { "Languages": "...", "Frameworks": "..." },
            "experience": [
                {
                    "role": "Job Title",
                    "company": "Company Name",
                    "duration": "Dates",
                    "location": "City",
                    "tech_stack": "Tools used",
                    "summary_input": "Full original text"
                }
            ]
        }
        """
        user_prompt = f"Parse this:\n{raw_text[:8000]}"
        response = self._call_llm(system_prompt, user_prompt)
        try:
            return ResumeData(**json.loads(response))
        except:
            return ResumeData()

    def audit_resume(self, resume_data: ResumeData, jd_text: str) -> ATSEvaluation:
        if not jd_text:
            return ATSEvaluation(score=100, missing_keywords=[], suggestions=["No JD provided for comparison."])

        # Use enhanced bullets if available, otherwise fallback to summary_input
        experience_text = []
        for job in resume_data.experience:
            if job.enhanced_bullets:
                experience_text.append(job.enhanced_bullets)
            else:
                experience_text.append(job.summary_input)

        resume_str = f"""
        Skills: {resume_data.skills}
        Experience: {experience_text}
        """

        system_prompt = """
        You are an ATS (Applicant Tracking System) Scanner. 
        Your job is to compare a Resume against a Job Description (JD).
        OUTPUT FORMAT (JSON):
        {
            "score": <integer 0-100>,
            "missing_keywords": ["keyword1", "keyword2"],
            "suggestions": ["suggestion1", "suggestion2"]
        }
        """
        user_prompt = f"JOB DESCRIPTION: {jd_text[:3000]}\nRESUME: {resume_str}"
        response_text = self._call_llm(system_prompt, user_prompt)
        try:
            return ATSEvaluation(**json.loads(response_text))
        except:
            return ATSEvaluation(score=0, missing_keywords=["Error parsing audit"], suggestions=[])

    def generate_cover_letter(self, resume_data: ResumeData, jd_text: str) -> str:
        resume_str = f"Name: {resume_data.full_name}\nExperience: {[job.enhanced_bullets for job in resume_data.experience]}"
        system_prompt = "Write a professional cover letter connecting the candidate's experience to the JD. Return JSON: { 'cover_letter': 'text...' }"
        user_prompt = f"Candidate: {resume_str}\nJD: {jd_text[:3000]}"
        response = self._call_llm(system_prompt, user_prompt)
        try:
            data = json.loads(response)
            return data.get("cover_letter", str(data))
        except:
            return response

    # --- UPDATED: REWRITE LOGIC WITH QUILLBOT EMULATION ---
    def rewrite_all_jobs(self, jobs: List[ExperienceItem], target_role: str, jd_text: str = "", tone: str = "Standard") -> List[ExperienceItem]:
        """
        Processes all jobs sequentially.
        Args:
            tone (str): "Standard" (Aggressive ATS) or "Humanized" (Quillbot Mode).
        """
        used_verbs: Set[str] = set()
        processed_jobs = []
        jd_context = jd_text[:3000] if jd_text else "General industry standards for this role."

        for job in jobs:
            forbidden_list_str = ", ".join(list(used_verbs))
            
            # --- TONE LOGIC ---
            if tone == "Humanized":
                # This prompt simulates Quillbot's "Humanize" logic
                style_guide = """
                MODE: QUILLBOT HUMANIZER / FLUENCY PARAPHRASER.
                Your goal is to rewrite the text to bypass AI detection and sound completely natural.
                
                STRICT GUIDELINES:
                1. SENTENCE VARIANCE: Do NOT start every sentence with a verb. Mix clauses (e.g., "By leveraging Python, I built..." instead of "Built...").
                2. VOCABULARY: Use simple, direct, high-fluency English. Avoid robotic words like 'Orchestrated', 'Pivotal', 'Spearheaded'.
                3. TONE: Write as if you are explaining your work to a friend, but professionally.
                4. STRUCTURE: Keep the metrics, but weave them naturally into the narrative.
                """
            else: # Standard / ATS Optimized
                style_guide = """
                MODE: HIGH-IMPACT ATS OPTIMIZER.
                1. START every bullet with a power verb (Engineered, Deployed, Architected).
                2. MAXIMIZE keyword density from the JD.
                3. BE AGGRESSIVE with metrics and results.
                """

            system_prompt = f"""
            You are an expert Resume Writer optimizing for a specific Job Description.
            TARGET ROLE: {target_role}
            
            {style_guide}
            
            JOB DESCRIPTION (JD) CONTEXT:
            {jd_context}
            
            CRITICAL RULES:
            1. UNIQUE ACTION VERBS: Do NOT use these words if possible: [{forbidden_list_str}].
            2. KEYWORD INJECTION: If the JD mentions specific skills matching the user's stack ({job.tech_stack}), include them.
            3. METRICS: Every bullet point must have a quantifiable result.
            """

            user_prompt = f"""
            Rewrite this job into 3 bullets.
            
            Role: {job.role}
            Tech Stack: {job.tech_stack}
            Raw Summary: {job.summary_input}
            
            Output format: JSON list of strings (e.g. {{ "bullets": ["..."] }})
            """

            response_text = self._call_llm(system_prompt, user_prompt)
            
            new_bullets = []
            try:
                data = json.loads(response_text)
                if isinstance(data, list): new_bullets = data
                elif "bullets" in data: new_bullets = data["bullets"]
                else: new_bullets = list(data.values())[0]
            except:
                new_bullets = [f"Managed {job.role} responsibilities.", "Optimized team workflows."]

            for bullet in new_bullets:
                words = bullet.split()
                if words:
                    first_word = words[0].strip(".,").capitalize()
                    used_verbs.add(first_word)
            
            job.enhanced_bullets = new_bullets
            processed_jobs.append(job)
            
        return processed_jobs

# --- DOCUMENT RENDERER (Layout Engine) ---
def create_styled_resume(data: ResumeData, filename="Generated_Resume.docx"):
    doc = Document()
    
    # 1. PAGE MARGINS (0.5 inch)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # 2. BASE FONT
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # 3. HELPER: BORDER FUNCTION
    def add_bottom_border(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bottom = OxmlElement('w:pBdr')
        border = OxmlElement('w:bottom')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), '000000')
        bottom.append(border)
        pPr.append(bottom)

    # 4. HEADER
    h_name = doc.add_paragraph()
    h_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h_name.add_run(data.full_name)
    run.bold = True
    run.font.size = Pt(16)
    
    h_contact = doc.add_paragraph()
    h_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_contact.add_run(data.contact_info)
    h_contact.paragraph_format.space_after = Pt(12)

    # 5. EDUCATION
    head_edu = doc.add_paragraph("EDUCATION")
    head_edu.runs[0].bold = True
    head_edu.runs[0].font.size = Pt(11)
    add_bottom_border(head_edu)

    for edu in data.education:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        r1 = p.add_run(edu.get('school', ''))
        r1.bold = True
        p.add_run(f"\t{edu.get('year', '')}")
        p2 = doc.add_paragraph(edu.get('degree', ''))
        p2.paragraph_format.space_after = Pt(8)

    # 6. SKILLS
    head_skills = doc.add_paragraph("SKILLS")
    head_skills.runs[0].bold = True
    head_skills.runs[0].font.size = Pt(11)
    add_bottom_border(head_skills)
    
    for cat, items in data.skills.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{cat}: ")
        r.bold = True
        p.add_run(str(items))
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 7. EXPERIENCE
    head_exp = doc.add_paragraph("WORK EXPERIENCE")
    head_exp.runs[0].bold = True
    head_exp.runs[0].font.size = Pt(11)
    add_bottom_border(head_exp)

    for job in data.experience:
        p_line1 = doc.add_paragraph()
        p_line1.paragraph_format.space_after = Pt(0)
        p_line1.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        
        r_role = p_line1.add_run(job.role)
        r_role.bold = True
        r_role.font.size = Pt(11)
        p_line1.add_run(f"\t{job.duration}")

        p_line2 = doc.add_paragraph()
        p_line2.paragraph_format.space_after = Pt(2)
        r_comp = p_line2.add_run(job.company)
        r_comp.italic = True
        p_line2.add_run(f" | {job.location}")

        for bullet in job.enhanced_bullets:
            p_bull = doc.add_paragraph(style='List Bullet')
            p_bull.add_run(bullet)
            p_bull.paragraph_format.space_after = Pt(0)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.save(filename)
    return filename